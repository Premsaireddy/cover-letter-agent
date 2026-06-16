import re
from io import BytesIO
<<<<<<< HEAD
=======

>>>>>>> 3739068cb7001a45761f34e25020b140149ccc75
import os
import streamlit as st
from dotenv import load_dotenv
from docx import Document
from openai import OpenAI
from pypdf import PdfReader

load_dotenv()
<<<<<<< HEAD

try:
    api_key = st.secrets["OPENAI_API_KEY"]
except (KeyError, FileNotFoundError):
    api_key = os.getenv("OPENAI_API_KEY")

client = OpenAI(api_key=api_key)
=======
>>>>>>> 3739068cb7001a45761f34e25020b140149ccc75

try:
    api_key = st.secrets["OPENAI_API_KEY"]
except (KeyError, FileNotFoundError):
    api_key = os.getenv("OPENAI_API_KEY")

client = OpenAI(api_key=api_key)
# ── Page config ────────────────────────────────────────────────────────────────
st.set_page_config(page_title="AI Cover Letter Agent", page_icon="📝")
st.title("AI Cover Letter Agent")
st.write(
    "Upload your CV, paste a job description, and the agent will autonomously "
    "decide how to draft, critique, and improve your cover letter."
)

uploaded_cv = st.file_uploader("Upload your CV", type=["pdf", "docx"])
job_description = st.text_area("Paste the job description", height=280)


# ── File extraction ─────────────────────────────────────────────────────────────

def extract_pdf_text(pdf_file):
    reader = PdfReader(pdf_file)
    return "\n".join(p.extract_text() for p in reader.pages if p.extract_text())


def extract_docx_text(docx_file):
    doc = Document(docx_file)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


# ── Content generation (LLM does the actual writing / critiquing) ───────────────
# These are NOT the agent — they are tools the agent calls to get work done.

def llm_generate(prompt):
    """Calls o3-mini for content generation tasks (writing, critiquing, improving)."""
    response = client.chat.completions.create(
        model="o3-mini",
        messages=[{"role": "user", "content": prompt}]
    )
    return response.choices[0].message.content


def _write_draft(cv_text, job_description):
    return llm_generate(f"""
You are an expert cover letter writer.
Write a tailored cover letter using the CV and job description below.

Rules:
- Do not invent experience. Use only information from the CV.
- Match the candidate's skills explicitly to the job requirements.
- Keep it professional and concise — between 250 and 400 words.
- Make it suitable for graduate, internship, or entry-level roles.
- Start with a specific hook referencing the company or role — never open with "I am writing to apply for..."
- Output only the cover letter text. No preamble, no labels, no commentary.

CV:
{cv_text}

Job Description:
{job_description}
""")


def _critique_draft(cv_text, job_description, draft):
    return llm_generate(f"""
You are a strict recruiter reviewing a cover letter.
Score it out of 10 by averaging these six criteria (each worth up to 1-2 points):
1. Relevance to the job description
2. Evidence drawn from the CV (concrete, not vague)
3. Professional tone
4. Specificity (names the company/role, references exact requirements)
5. No invented claims (everything traceable to the CV)
6. Clear structure (strong opening, body, closing)

YOU MUST respond in EXACTLY this format — no preamble, no variation whatsoever:
Score: X/10

Feedback:
- point 1
- point 2
- point 3

CV:
{cv_text}

Job Description:
{job_description}

Cover Letter:
{draft}
""")


def _improve_draft(cv_text, job_description, draft, critique, score, target_score):
    return llm_generate(f"""
You are an expert cover letter writer.
The previous cover letter scored {score}/10. The target is {target_score}/10.
Improve it by carefully addressing every point in the recruiter's feedback below.

Rules:
- Do not invent anything not present in the CV or job description.
- Keep it between 250 and 400 words.
- Be more specific: reference the company name, exact role title, and concrete requirements from the job description.
- Fix every issue raised in the feedback — do not ignore any point.
- Keep the strong parts of the original letter intact.
- Output only the improved cover letter text. No preamble, no labels, no commentary.

CV:
{cv_text}

Job Description:
{job_description}

Original Cover Letter:
{draft}

Recruiter Feedback:
{critique}
""")


def extract_score(critique_text):
    match = re.search(r"Score:\s*(\d+(?:\.\d+)?)\s*/\s*10", critique_text, re.IGNORECASE)
    return float(match.group(1)) if match else 0.0


# ── Tool schema (what the agent sees and reasons over) ─────────────────────────
# The agent (gpt-4o) reads these descriptions and DECIDES which tool to call next.
# This is the core of what makes it agentic — the LLM is the decision-maker.

AGENT_TOOLS = [
    {
        "type": "function",
        "function": {
            "name": "write_draft",
            "description": (
                "Write the very first draft of the cover letter from the CV and job description. "
                "Always call this tool first before anything else."
            ),
            "parameters": {"type": "object", "properties": {}, "required": []}
        }
    },
    {
        "type": "function",
        "function": {
            "name": "critique_draft",
            "description": (
                "Score and critique the current cover letter draft out of 10 with specific feedback. "
                "Call this after writing or improving a draft to evaluate quality."
            ),
            "parameters": {"type": "object", "properties": {}, "required": []}
        }
    },
    {
        "type": "function",
        "function": {
            "name": "improve_draft",
            "description": (
                "Revise and improve the current draft based on the latest critique feedback. "
                "Only call this if the current score is below the target (8/10) "
                "AND fewer than 3 improvement attempts have been made."
            ),
            "parameters": {"type": "object", "properties": {}, "required": []}
        }
    },
    {
        "type": "function",
        "function": {
            "name": "finalize",
            "description": (
                "Mark the cover letter as complete and ready to submit. "
                "Call this when the score reaches 8/10 or above, "
                "OR when 3 improvement attempts have been exhausted."
            ),
            "parameters": {"type": "object", "properties": {}, "required": []}
        }
    }
]


# ── Tool executor (runs the actual work when the agent picks a tool) ────────────

def execute_tool(tool_name, state, status_box):
    """
    Executes whichever tool the agent decided to call.
    Updates shared state and returns a result string fed back to the agent.
    """
    if tool_name == "write_draft":
        status_box.info("✍️ Agent → write_draft: Writing first draft...")
        draft = _write_draft(state["cv_text"], state["job_description"])
        state["current_draft"] = draft
        return f"Draft written. Word count: {len(draft.split())} words."

    elif tool_name == "critique_draft":
        status_box.info("🕵️ Agent → critique_draft: Evaluating current draft...")
        critique = _critique_draft(
            state["cv_text"], state["job_description"], state["current_draft"]
        )
        score = extract_score(critique)
        state["current_critique"] = critique
        state["current_score"] = score
        state["attempts_log"].append({
            "attempt": len(state["attempts_log"]) + 1,
            "score": score,
            "critique": critique
        })
        return (
            f"Critique complete. Score: {score}/10. "
            f"Improvement attempts used so far: {state['improvement_attempts']}/{state['max_attempts']}."
        )

    elif tool_name == "improve_draft":
        state["improvement_attempts"] += 1
        status_box.warning(
            f"⚠️ Agent → improve_draft: Score was {state['current_score']}/10. "
            f"Revising (attempt {state['improvement_attempts']}/{state['max_attempts']})..."
        )
        improved = _improve_draft(
            state["cv_text"],
            state["job_description"],
            state["current_draft"],
            state["current_critique"],
            state["current_score"],
            state["target_score"]
        )
        state["current_draft"] = improved
        return (
            f"Draft improved. "
            f"{state['improvement_attempts']} of {state['max_attempts']} attempts used."
        )

    elif tool_name == "finalize":
        state["finalized"] = True
        if state["current_score"] >= state["target_score"]:
            status_box.success(
                f"🎉 Agent → finalize: Target reached ({state['current_score']}/10). Cover letter ready."
            )
        else:
            status_box.error(
                f"🛑 Agent → finalize: Max attempts exhausted. Best score: {state['current_score']}/10."
            )
        return "Cover letter finalized and ready for download."

    return f"Unknown tool called: {tool_name}"


# ── Agent loop (THIS is the agentic part) ──────────────────────────────────────
# gpt-4o acts as the agent brain. It reads the tool descriptions + conversation
# history and DECIDES which tool to call next. There are no Python if-statements
# controlling the flow — the LLM is the decision-maker at every step.

def run_agent(cv_text, job_description, status_box):
    state = {
        "cv_text": cv_text,
        "job_description": job_description,
        "current_draft": "",
        "current_critique": "",
        "current_score": 0.0,
        "improvement_attempts": 0,
        "max_attempts": 3,
        "target_score": 8,
        "attempts_log": [],
        "finalized": False
    }

    # System prompt gives the agent its goal and decision rules.
    # Crucially, the agent REASONS about these — they are not enforced by code.
    system_prompt = f"""
You are an autonomous Cover Letter Agent. Your sole goal is to produce the best
possible cover letter for a job applicant by reasoning and using your tools.

Tools available to you:
- write_draft     → Always call this first to produce an initial draft.
- critique_draft  → Call after writing or improving to get a score and feedback.
- improve_draft   → Call if the score is below {state['target_score']}/10 and you have used fewer than {state['max_attempts']} improvement attempts.
- finalize        → Call when the score reaches {state['target_score']}/10 or above, OR when {state['max_attempts']} improvement attempts have been exhausted.

You reason about which tool to call at every step. You are the decision-maker.
The tool result will always tell you the current score and how many attempts remain.
"""

    messages = [
        {"role": "system", "content": system_prompt},
        {
            "role": "user",
            "content": (
                "Begin. Autonomously produce the best possible cover letter "
                "for this candidate using your tools."
            )
        }
    ]

    safety_limit = 12  # Prevents runaway loops in edge cases

    for _ in range(safety_limit):
        if state["finalized"]:
            break

        # The agent (gpt-4o) looks at the full conversation + tool results so far
        # and decides on its own which tool to call next.
        response = client.chat.completions.create(
            model="o3-mini",          # Agent brain: reasons and picks tools
            messages=messages,
            tools=AGENT_TOOLS,
            tool_choice="required"   # Agent must always pick a tool (keeps loop clean)
        )

        agent_message = response.choices[0].message
        messages.append(agent_message)

        if not agent_message.tool_calls:
            break  # Safety exit (shouldn't happen with tool_choice="required")

        # Execute every tool the agent chose and feed results back into the conversation
        for tool_call in agent_message.tool_calls:
            result = execute_tool(tool_call.function.name, state, status_box)
            messages.append({
                "role": "tool",
                "tool_call_id": tool_call.id,
                "content": result
            })

    return state


# ── DOCX export ────────────────────────────────────────────────────────────────

def create_docx(text):
    doc = Document()
    doc.add_heading("Cover Letter", level=1)
    for paragraph in text.split("\n"):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ── UI trigger ─────────────────────────────────────────────────────────────────

if st.button("Generate Cover Letter"):
    if uploaded_cv is None:
        st.error("Please upload your CV as PDF or DOCX.")
        st.stop()
    if not job_description.strip():
        st.error("Please paste the job description.")
        st.stop()

    with st.spinner("Reading your CV..."):
        if uploaded_cv.name.lower().endswith(".pdf"):
            cv_text = extract_pdf_text(uploaded_cv)
        elif uploaded_cv.name.lower().endswith(".docx"):
            cv_text = extract_docx_text(uploaded_cv)
        else:
            st.error("Unsupported file type.")
            st.stop()

    if not cv_text.strip():
        st.error("Could not read your CV. Please check the file.")
        st.stop()

    status_box = st.empty()
    final_state = run_agent(cv_text, job_description, status_box)

    st.session_state["final_letter"] = final_state["current_draft"]
    st.session_state["attempts_log"] = final_state["attempts_log"]


# ── Results ────────────────────────────────────────────────────────────────────

if "final_letter" in st.session_state:
    st.subheader("Agent Decision Log")
    for item in st.session_state["attempts_log"]:
        with st.expander(f"Attempt {item['attempt']} — Score: {item['score']}/10"):
            st.write(item["critique"])

    st.subheader("Final Cover Letter")
    st.write(st.session_state["final_letter"])

    st.download_button(
        label="Download Cover Letter as DOCX",
        data=create_docx(st.session_state["final_letter"]),
        file_name="cover_letter.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
